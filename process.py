import azure.cognitiveservices.speech as speechsdk
import time
# import process_sentences
from collections import defaultdict
from nltk.tokenize import word_tokenize, sent_tokenize
import matplotlib.pyplot as plt
import docx
import sys
import re
mysp=__import__("my-voice-analysis")

debug = False
OUTPUT_FILE_PATH = "speechReport.docx"

def debug_print(*args):
    if(debug):
        print(*args)

with open("input.txt", "r") as file:
    input_speech= file.read()

tokens_input = word_tokenize(input_speech)
word_counts_input = {}
word_counts_input = word_counts_input.fromkeys(tokens_input)
for i in word_counts_input:
    word_counts_input[i] = tokens_input.count(i)
# debug_print("word_counts_input: ", word_counts_input)


filler_words = [ "um", "ah", "basically", "umm", "right", "actually", "so", "like" ]
profanity = ["shit", "damn", "fuck", "heck", "hell", "bullshit", "crap", "nigga"]


def speech_recognize_continuous_from_file(file):
    """performs continuous speech recognition with input from an audio file"""
    # <SpeechContinuousRecognitionWithFile>
    speech_config = speechsdk.SpeechConfig(subscription="75ea7877e1ac405aac11036d4fe62f37", region="eastus")
    audio_config = speechsdk.audio.AudioConfig(filename=file)

    speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)

    done = False
    text = ""

    def stop_cb(evt):
        """callback that signals to stop continuous recognition upon receiving an event `evt`"""
        nonlocal text
        nonlocal done
        if(evt.result.reason == speechsdk.ResultReason.RecognizedSpeech):
            text += evt.result.text + "\n\n"
        if(evt.result.reason == speechsdk.ResultReason.Canceled):
            done = True
        
    speech_recognizer.recognized.connect(stop_cb)
    speech_recognizer.canceled.connect(stop_cb)

    # Start continuous speech recognition
    speech_recognizer.start_continuous_recognition()
    while not done:
        time.sleep(.5)

    speech_recognizer.stop_continuous_recognition()
    return text

def process_sentences_filler_words_and_profanity(text):
    global word_counts_text
    profanity_words_detected = []
    filler_words_detected = []
    for i in profanity:
        if i in word_counts_text:
            profanity_words_detected.append({i: word_counts_text[i]})

    for i in filler_words:
        if(i in word_counts_text):
            filler_words_detected.append({i: word_counts_text[i]})

    return filler_words_detected, profanity_words_detected

def process_tone(text):
    feedback = {}

    global word_counts_input
    global word_counts_text
    num_paras_actual = input_speech.count("\n\n")
    num_paras_transcribed = text.count("\n\n")

    debug_print("num_paras_actual: ", num_paras_actual)
    debug_print("num_paras_transcribed: ", num_paras_transcribed)

    if(num_paras_actual < num_paras_transcribed):
        feedback["paras"] = 1
    elif(num_paras_actual > num_paras_transcribed):
        feedback["paras"] = -1
    else:
        feedback["paras"] = 0
    try:
        num_sentences_actual = word_counts_input["."]
    except:
        num_sentences_actual = 0

    try:
        num_sentences_transcribed = word_counts_text["."]
    except:
        num_sentences_transcribed = 0

    debug_print("num_sentences_actual: ", num_sentences_actual)
    debug_print("num_sentences_transcribed: ", num_sentences_transcribed)

    if(num_sentences_actual < num_sentences_transcribed):
        feedback["sentences"] = 1
    elif(num_sentences_actual > num_sentences_transcribed):
        feedback["sentences"] = -1
    else:
        feedback["sentences"] = 0


    try: 
        num_questions_actual = word_counts_input["?"]
    except:
        num_questions_actual = 0
    
    try:
        num_questions_transcribed = word_counts_text["?"]
    except:
        num_questions_transcribed = 0

    debug_print("num_questions_actual: ", num_questions_actual)
    debug_print("num_questions_transcribed: ", num_questions_transcribed)

    if(num_questions_actual < num_questions_transcribed): 
        feedback["questions"] = 1
    elif(num_questions_actual > num_questions_transcribed):
        feedback["questions"] = -1
    else:
        feedback["questions"] = 0

    try:
        num_exclamations_actual = word_counts_input["!"]
    except:
        num_exclamations_actual = 0
    try:
        num_exclamtions_transcribed = word_counts_text["!"]
    except:
        num_exclamtions_transcribed = 0

    debug_print("num_exclamations_actual: ", num_exclamations_actual)
    debug_print("num_exclamtions_transcribed: ", num_exclamtions_transcribed)

    if(num_exclamations_actual < num_exclamtions_transcribed):
        feedback["exclamation"] = 1
    
    elif(num_exclamations_actual > num_exclamtions_transcribed):
        feedback["exclamation"] = -1
    
    else:
        feedback["exclamation"] = 0


    try:
        num_commas_actual = word_counts_input[","]
    except:
        num_commas_actual = 0
    try:
        num_commas_transcribed = word_counts_text[","]
    except:
        num_commas_transcribed = 0

    debug_print("num_commas_actual: ", num_commas_actual)
    debug_print("num_commas_transcribed: ", num_commas_transcribed)

    if(num_commas_actual < num_commas_transcribed):
        feedback["commas"] = 1
    
    elif(num_commas_actual > num_commas_transcribed):
        feedback["commas"] = -1
    
    else:
        feedback["commas"] = 0

    return feedback

text = speech_recognize_continuous_from_file("temp_audio.wav")
debug_print("Text Transcribed from Speech Service")
debug_print("*" * 100)
debug_print("transcribed Text: ", text)
with open("transcribed_text.txt", "w") as file:
    file.write(text)
debug_print("*" * 100)
debug_print("\n" * 5)


# # with open("transcribed_text.txt", "r") as file:
# #     text = file.read()

text = text.lower()
tokens_text = word_tokenize(text)
word_counts_text = {}
word_counts_text = word_counts_text.fromkeys(tokens_text)
for i in word_counts_text:
    word_counts_text[i] = tokens_text.count(i)
debug_print("word_counts_text: ", word_counts_text)

tone_feedback = process_tone(text)
debug_print("feedback: ", tone_feedback)


filler_words_detected, profanity_words_detected  = process_sentences_filler_words_and_profanity(text)
debug_print("filler_words_detected: ", filler_words_detected)
debug_print("profanity_words_detected: ", profanity_words_detected)


def generateFeeback(tone_feedback, filler_words_detected, profanity_words_detected):
    document = docx.Document()
    global word_counts_text
    document.add_heading("Speech Report", 0)
    document.add_paragraph("A detailed report I could come up with while listening to you speak! Hope it could be of some help.")
    # document.add_heading("Analysis of your tone of speech", 2)
    tone_feedback_message = {
        "paras": {
            -1: "It appears that the deliberate long pauses between the different paragraphs could be something we could improve upon.",
            0: "Spot on with the deliberate long pauses between the different paragraphs.",
            1: "Seems like there were quite a few extra long pauses you had right there. Maybe we could improve upon exerting those long pauses as approproate, and not overdoing it. :)"
        },
        "sentences": {
            -1: "We may have hurried up a bit there. Felt like some sentences couldn't be distinguished from the other. If we could work upon those tiny little pauses to be able to differentiate sentences, may be the audience could connect better.",
            0: "Seems like every sentence you delivered was crisp and clear. Way to go!",
            1: "We might have had a few unnecessary pauses mid-sentence in some places. Let's work on delivering each sentence, clearly and confidently without breaks!"
        },
        "questions": {
            -1: "Did we miss a question out there. Seems like it didnt sound like one. Maybe if we could get the right tone to stimulate a train of thoughts in the audience, to make them curious for an answer, it would be more engaging. Don't you think?", 
            0: "Every question you put out there was well said. Good job! Lets keep the tone rolling!",
            1: "Did we just end up making something that's not a queston, sound like one and leave the audience confused? Let's work on it! We could do better."
        },
        "exclamation": {
            -1: "Seems like there is a bit of excitement lacking. Bring it out. Seems like we need to go Hurraayy!!!! a bit more. The secret to getting a listener excited and wanting for more is the energy that you radiate. Let's get the energy up!",
            0: "You nailed all the exclamations with all the energy you got! Way to go!",
            1: "Ooh! That was a high energy exciting talk.",
        },
        "commas": {
            -1: "We could probably work on those soft pauses during a sentence.",
            0: "The soft pauses during the sentences capturing the attention of the audience done graciously. Job Well done!!",
            1: "Seemed like we had one too many soft pauses in a sentence right there. Let us work on getting them right!" 
        }
    }

    document.add_paragraph(tone_feedback_message["paras"][tone_feedback["paras"]], style='List Bullet')
    document.add_paragraph(tone_feedback_message["sentences"][tone_feedback["sentences"]], style='List Bullet')
    document.add_paragraph(tone_feedback_message["questions"][tone_feedback["questions"]], style='List Bullet')
    document.add_paragraph(tone_feedback_message["exclamation"][tone_feedback["exclamation"]], style='List Bullet')
    document.add_paragraph(tone_feedback_message["commas"][tone_feedback["commas"]], style='List Bullet')


    if(len(filler_words_detected) > 0):
        filler_words_detected.sort(key=lambda x: list(x.values())[0], reverse=True)
        document.add_heading("Filler Words details", 2)
        document.add_paragraph("Your most used filler words are: ")
        for i in range(min(5, len(filler_words_detected))):
            item = list(filler_words_detected[i].items())[0]
            document.add_paragraph(f'{item[0]} used {item[1]} times.', style='List Bullet')

        total_words_spoken = sum(list(word_counts_text.values()))
        # total_filler_words_spoken = sum([i for i in list(filler_words_detected[i].values())[0]])
        total_filler_words_spoken = sum(list(map(lambda x: list(x.values())[0], filler_words_detected)))

        percentage_filler_words = (total_filler_words_spoken / total_words_spoken) * 100
        document.add_paragraph(f"We have another interesting insight that might help you. Looks like your speech contains {round(percentage_filler_words, 2)} of filler words! Hope this helps you improve refining your talk")

    if(len(profanity_words_detected) > 0):
        document.add_heading("Information about profanity", 2)
        document.add_paragraph("Words that you might have uttered unconsciously that would probably need a little reconsideration")
        profanity_words_detected.sort(key=lambda x: list(x.values())[0], reverse=True)
        for i in range(len(profanity_words_detected)):
            item = list(profanity_words_detected[i].items())[0]
            document.add_paragraph(f'{item[0]} used {item[1]} times.', style='List Bullet')

    stats_for_nerds_table = get_stats_for_nerds("temp_audio", ".")

    table = document.add_table(rows=0, cols=2)
    for key, value in stats_for_nerds_table:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value

    document.save(OUTPUT_FILE_PATH)
    return OUTPUT_FILE_PATH


def get_stats_for_nerds(audio_filename, audio_file_directory):
    # print("\nStats for Nerds :P\n")
    sys.stdout = open("stats_for_nerds.txt", "w")
    mysp.mysptotal(audio_filename, audio_file_directory)
    sys.stdout = sys.__stdout__
    stats_file = open("stats_for_nerds.txt", "r")
    stats = stats_file.read().splitlines()[2:]
    legend_titles = ['No. of syllables in your speech', 'No. of pauses in your speech', 
                     'Average syllables per second', 'Average syllables per second (pause time excluded)',
                     'Speech duration in seconds (pause time excluded)', 'Speech duration in seconds',
                     'Ratio of speaking and total duration', 'Fundamental frequency distribution mean',
                     'Fundamental frequency distribution std', 'Fundamental frequency distribution median',
                     'Global min. of Fundamental frequency distribution', 'Global max. of Fundamental frequency distribution', 
                     'Global 25th percentile of FFD', 'Global 75th percentile of FFD']

    for index in range(len(legend_titles)):
        stats[index][0] = legend_titles[0]
    return stats


feedback_report = generateFeeback(tone_feedback, filler_words_detected, profanity_words_detected)
# print(feedback_report)
# stats_for_nerds("temp_audio", ".")